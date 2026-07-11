from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("reports") / "rapport_court_projet11.docx"
AUTHOR = "Vina RAHARITSIFA - M1 I2AD, INSI"


sections = [
    (
        "Introduction",
        [
            "Ce projet porte sur l'analyse automatique du sentiment dans des discours politiques publies sur Twitter. L'objectif est de classer un texte court en trois categories, negative, neutre ou positive, afin de fournir un outil simple d'aide a l'analyse de tendances discursives. Le contexte politique est particulierement interessant pour le traitement automatique du langage naturel, car les messages sont courts, souvent polarises, et riches en hashtags, mentions, slogans ou references a des evenements.",
            "Le travail a ete realise en plusieurs etapes : exploration des donnees, nettoyage, labellisation automatique, vectorisation, entrainement de plusieurs modeles, suivi des experiences avec MLflow, sauvegarde du meilleur modele, puis deploiement sous forme d'API FastAPI et de dashboard Streamlit. L'entrainement complet a ete effectue sur Google Colab afin de disposer d'un environnement adapte au volume de donnees et de faciliter la sauvegarde des artefacts dans Google Drive.",
        ],
    ),
    (
        "Jeux de donnees",
        [
            "Trois sources Kaggle ont ete utilisees. Le premier dataset regroupe des tweets politiques generaux collectes autour du hashtag Politics. Le deuxieme concerne le conflit en Ukraine et rassemble des messages associes a plusieurs hashtags. Le troisieme porte sur l'election presidentielle americaine de 2020, avec des tweets lies a Donald Trump et Joe Biden. Ces trois sources permettent de couvrir plusieurs contextes politiques, plusieurs vocabulaires et plusieurs niveaux de polarisation.",
            "Apres nettoyage et deduplication, le pipeline de preprocessing a produit 1 695 833 tweets exploitables. Pour la phase de modelisation, un echantillonnage stratifie a ete applique afin de limiter le desequilibre et de garder un volume raisonnable pour comparer les modeles. Le jeu final utilise pour l'entrainement contient 600 000 tweets, repartis en trois classes de 200 000 exemples chacune. Le split train, validation et test est de 420 000, 90 000 et 90 000 tweets.",
        ],
    ),
    (
        "Preprocessing NLP",
        [
            "Le preprocessing a ete adapte aux particularites de Twitter. Les URLs, mentions, retweets et caracteres non alphabetiques sont supprimes. Les hashtags sont conserves sous forme de mots, car ils portent souvent une information semantique importante. Le texte est ensuite converti en minuscules, tokenise, filtre avec une liste de stopwords et lemmatise avec NLTK. Certains mots politiques comme vote, election, trump, biden, ukraine, russia, nato, war ou peace sont volontairement conserves meme s'ils peuvent apparaitre dans des listes de mots frequents.",
            "La labellisation des sentiments a ete realisee automatiquement avec VADER. Cette approche est coherente pour des textes courts issus des reseaux sociaux, meme si elle introduit une limite importante : les labels ne proviennent pas d'une annotation humaine experte. Le modele apprend donc a reproduire une approximation lexicale du sentiment, et non une verite absolue. Cette limite doit etre prise en compte dans l'interpretation des resultats.",
        ],
    ),
    (
        "Modelisation",
        [
            "La representation textuelle retenue est TF-IDF avec unigrammes et bigrammes, limitee aux 50 000 features les plus utiles. Les bigrammes permettent de mieux conserver certaines expressions courtes, notamment les negations comme not good ou les associations politiques frequentes. Cinq classifieurs ont ete compares : Logistic Regression, Naive Bayes, LinearSVC, Random Forest et Gradient Boosting.",
            "Le meilleur resultat est obtenu par LinearSVC, avec un F1 test de 0.8902 et une accuracy test de 0.8904. Logistic Regression obtient egalement un bon score, avec un F1 test de 0.8821. Les modeles Naive Bayes, Random Forest et Gradient Boosting sont moins performants sur ce jeu de donnees. Le choix final de LinearSVC est donc justifie par son compromis entre performance, stabilite et vitesse d'inference.",
        ],
    ),
    (
        "Suivi MLflow et artefacts",
        [
            "Les experiences de modelisation sont suivies avec MLflow. Les metriques principales conservees sont l'accuracy, le F1-score, les scores de validation croisee et le temps d'entrainement. Les artefacts sauvegardes incluent le meilleur modele, le vectoriseur TF-IDF, les rapports de classification et les graphiques comparatifs. Cette organisation facilite la verification des resultats et la reproductibilite du projet.",
            "Le modele final est sauvegarde dans models/best_model.pkl et le vectoriseur dans models/tfidf_vectorizer.pkl. Ces deux fichiers suffisent pour effectuer une prediction en production, sans charger les datasets bruts. Le dashboard Streamlit inclut aussi un mecanisme de telechargement automatique des artefacts depuis Google Drive lorsque les fichiers ne sont pas presents localement.",
        ],
    ),
    (
        "Deploiement",
        [
            "Le projet expose deux interfaces. La premiere est une API FastAPI avec les routes /health, /predict et /predict/batch. Elle permet d'integrer le modele dans une application externe ou de tester les predictions depuis la documentation interactive /docs. La deuxieme interface est un dashboard Streamlit, deja deploye, qui presente les resultats de modelisation, les statistiques des datasets, des textes de demonstration et une analyse en temps reel.",
            "Docker Compose permet de lancer l'API, Streamlit et MLflow dans un environnement local coherent. Les fichiers Dockerfile.api, Dockerfile.streamlit, docker-compose.yml et docker-compose.hub.yml separent le build local du lancement via images Docker Hub. Cette separation rend le projet plus lisible pour un rendu academique et plus simple a tester sur une autre machine.",
        ],
    ),
    (
        "Limites et perspectives",
        [
            "La principale limite concerne la labellisation automatique. VADER est utile pour amorcer un projet de sentiment analysis, mais il peut mal gerer l'ironie, les references implicites, le sarcasme ou certaines formulations politiques ambigues. Une perspective importante serait donc de constituer un sous-ensemble annote manuellement pour evaluer plus finement la qualite reelle des predictions.",
            "Une autre amelioration possible serait de comparer LinearSVC avec des modeles de langue modernes, par exemple des embeddings contextualises ou un modele Transformer fine-tune sur des tweets politiques. Ces approches pourraient mieux capturer le contexte, mais elles demandent davantage de ressources et une procedure d'evaluation plus stricte. Pour le cadre de ce projet, la solution TF-IDF et LinearSVC reste pertinente, rapide et explicable.",
        ],
    ),
    (
        "Conclusion",
        [
            "Le projet aboutit a une chaine complete de NLP appliquee : donnees explorees et nettoyees, modele entraine sur Colab, suivi MLflow, artefacts sauvegardes, API fonctionnelle, dashboard Streamlit et tests pytest. Les resultats montrent que LinearSVC est le meilleur modele parmi ceux compares, avec un F1 test de 0.8902. Le rendu respecte ainsi les livrables attendus tout en gardant une architecture simple et exploitable.",
        ],
    ),
]


def set_cell_text(cell, text, bold=False):
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")


def style_paragraph(paragraph, size=11, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    paragraph.alignment = align
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.bold = bold
        run.font.color.rgb = RGBColor(0, 0, 0)
        run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "000000")
        borders.append(tag)
    tbl_pr.append(borders)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header.is_linked_to_previous = False
section.footer.is_linked_to_previous = False

styles = doc.styles
for name in ("Normal", "Heading 1", "Heading 2"):
    style = styles[name]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11 if name == "Normal" else 12)
    style.font.color.rgb = RGBColor(0, 0, 0)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Projet 11 - Analyse des discours politiques")
run.bold = True
style_paragraph(title, size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.add_run(f"Auteur : {AUTHOR}")
style_paragraph(subtitle, align=WD_ALIGN_PARAGRAPH.CENTER)

meta = doc.add_paragraph()
meta.add_run("Rapport court de projet NLP. Entrainement realise sur Google Colab, modele sauvegarde dans models/, API FastAPI et dashboard Streamlit disponibles dans le depot.")
style_paragraph(meta)

doc.add_paragraph()

table = doc.add_table(rows=1, cols=4)
set_table_borders(table)
headers = ["Modele", "F1 test", "Accuracy test", "CV F1"]
for cell, header in zip(table.rows[0].cells, headers):
    set_cell_text(cell, header, bold=True)
rows = [
    ("LinearSVC", "0.8902", "0.8904", "0.8835"),
    ("Logistic Regression", "0.8821", "0.8824", "0.8739"),
    ("Naive Bayes", "0.7224", "0.7226", "0.7214"),
    ("Random Forest", "0.6994", "0.7021", "0.6962"),
    ("Gradient Boosting", "0.6168", "0.6171", "0.6153"),
]
for values in rows:
    cells = table.add_row().cells
    for cell, value in zip(cells, values):
        set_cell_text(cell, value)

doc.add_paragraph()

for heading, paragraphs in sections:
    h = doc.add_paragraph()
    h.add_run(heading).bold = True
    style_paragraph(h, size=12, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    for paragraph_text in paragraphs:
        p = doc.add_paragraph(paragraph_text)
        style_paragraph(p)

doc.save(OUT)
print(OUT)
